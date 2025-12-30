import requests
import pandas as pd
import os
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

# ================= CONFIG =================
SONAR_URL = "http://tu-sonarqube"
SONAR_TOKEN = "TU_TOKEN_DE_ACCESO"
AUTH = (SONAR_TOKEN, "")

OUTPUT_BASE = "reports"

PROJECT_KEYS = [
    # puedes agregar TODOS los proyectos aquí
]

os.makedirs(OUTPUT_BASE, exist_ok=True)

# ================= API =================
def get_issues(project):
    issues_all = []
    page = 1

    while True:
        r = requests.get(
            f"{SONAR_URL}/api/issues/search",
            params={
                "componentKeys": project,
                "types": "VULNERABILITY,BUG,CODE_SMELL",
                "ps": 500,
                "p": page
            },
            auth=AUTH
        )

        if r.status_code != 200:
            break

        data = r.json()
        issues = data.get("issues", [])
        issues_all.extend(issues)

        if len(issues) < 500:
            break

        page += 1

    return issues_all


def get_coverage(project):
    r = requests.get(
        f"{SONAR_URL}/api/measures/component",
        params={"component": project, "metricKeys": "coverage"},
        auth=AUTH
    )

    if r.status_code != 200:
        return "N/A"

    measures = r.json().get("component", {}).get("measures", [])
    return measures[0]["value"] if measures else "N/A"

# ================= CHARTS =================
def generate_charts(df, chart_dir):
    charts = {}

    def save(series, title, name):
        plt.figure(figsize=(8,4))
        series.plot(kind="bar")
        plt.title(title)
        plt.ylabel("Cantidad")
        plt.tight_layout()
        path = os.path.join(chart_dir, name)
        plt.savefig(path)
        plt.close()
        return path

    charts["severidad"] = save(
        df["Severidad"].value_counts(),
        "Issues por Severidad",
        "severidad.png"
    )

    charts["tipo"] = save(
        df["Tipo"].value_counts(),
        "Issues por Tipo",
        "tipo.png"
    )

    charts["estado"] = save(
        df["Estado"].value_counts(),
        "Issues por Estado",
        "estado.png"
    )

    return charts

# ================= REPORTES =================
def generate_word(project, df, summary, coverage, charts, base_path):
    doc = Document()
    doc.add_heading(f"Reporte de Seguridad – {project}", 0)

    doc.add_paragraph("SonarQube Community Edition")
    doc.add_paragraph("Pipeline: Azure DevOps")
    doc.add_paragraph(f"Cobertura de Pruebas: {coverage}%")

    doc.add_heading("Análisis Gráfico", level=1)
    for img in charts.values():
        doc.add_picture(img, width=Cm(14))

    doc.add_heading("Resumen Ejecutivo", level=1)
    for _, r in summary.iterrows():
        doc.add_paragraph(
            f"{r['Tipo']} | {r['Severidad']}: {r['Cantidad']}"
        )

    doc.add_page_break()
    doc.add_heading("Detalle de Hallazgos", level=1)

    for _, r in df.iterrows():
        doc.add_paragraph(
            f"{r['Tipo']} ({r['Severidad']})\n"
            f"Regla: {r['Regla']}\n"
            f"Archivo: {r['Archivo']}\n"
            f"{r['Mensaje']}\n"
            f"Estado: {r['Estado']}\n"
        )

    doc.save(os.path.join(base_path, f"{project}.docx"))


def generate_pdf(project, summary, coverage, charts, base_path):
    pdf = SimpleDocTemplate(
        os.path.join(base_path, f"{project}.pdf"),
        pagesize=A4
    )

    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(f"Reporte de Seguridad – {project}", styles["Title"]))
    story.append(Spacer(1, 0.7*cm))
    story.append(Paragraph(f"Cobertura de Pruebas: {coverage}%", styles["Normal"]))

    story.append(Spacer(1, 0.7*cm))
    story.append(Paragraph("Análisis Gráfico", styles["Heading2"]))

    for img in charts.values():
        story.append(Image(img, width=14*cm, height=7*cm))
        story.append(Spacer(1, 0.4*cm))

    table_data = [["Tipo", "Severidad", "Cantidad"]]
    for _, r in summary.iterrows():
        table_data.append([r["Tipo"], r["Severidad"], str(r["Cantidad"])])

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Resumen Ejecutivo", styles["Heading2"]))
    story.append(Table(table_data, colWidths=[5*cm, 5*cm, 4*cm]))

    pdf.build(story)

# ================= MAIN =================
for project in PROJECT_KEYS:
    print(f"🔍 Procesando {project}")

    base_path = os.path.join(OUTPUT_BASE, project)
    chart_dir = os.path.join(base_path, "charts")

    os.makedirs(chart_dir, exist_ok=True)

    issues = get_issues(project)

    if not issues:
        print(f"⚠ Sin issues: {project}")
        continue

    rows = [{
        "Tipo": i["type"],
        "Severidad": i["severity"],
        "Regla": i["rule"],
        "Mensaje": i["message"],
        "Archivo": i.get("component", ""),
        "Estado": i["status"]
    } for i in issues]

    df = pd.DataFrame(rows)

    summary = (
        df.groupby(["Tipo", "Severidad"])
          .size()
          .reset_index(name="Cantidad")
    )

    coverage = get_coverage(project)
    charts = generate_charts(df, chart_dir)

    generate_word(project, df, summary, coverage, charts, base_path)
    generate_pdf(project, summary, coverage, charts, base_path)

print("✅ Reportes por proyecto generados correctamente")
